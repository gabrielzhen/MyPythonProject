from docx import Document
from pathlib import Path

path=Path('F:\\编程文件\\拆分与合并\\子文件夹')
doc2=Document()
for file in path.rglob('*.docx'):
    doc=Document(file)
    for paragraph in doc.paragraphs:
        print(paragraph.text)
        print('++++++++++++++')
        newpara=doc2.add_paragraph()
        run=newpara.add_run(paragraph.text)
        run.bold=True
        run.underline=True
doc2.save('F:\\编程文件\\拆分与合并\\合并.docx')